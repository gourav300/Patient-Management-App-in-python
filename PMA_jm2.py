from tkinter import *
from tkinter import font as tkFont
from tkinter import ttk
from tkinter import messagebox
import xlsxwriter
from docx import Document
import os
import datetime
import sqlite3
import shutil
import subprocess
import openpyxl
from openpyxl import workbook
#import pandas as pd
#from openpyxl.drawing.image import Image
import webbrowser
#import PIL


date = datetime.datetime.now().date().strftime('%d-%m-%Y')
date = str(date)


class Application(object):
    def __init__(self, master):

        self.parent_dir = os.getcwd()

        self.name = StringVar()
        self.disease = StringVar()
        self.age = StringVar()
        self.mobile = StringVar()
        self.gender = StringVar()

        self.name_update = StringVar()
        self.disease_update = StringVar()
        self.age_update = StringVar()
        self.mobile_update = StringVar()
        self.gender_update = StringVar()

        self.con = sqlite3.connect('database.db')
        self.cur = self.con.cursor()

        data_table = """CREATE TABLE IF NOT EXISTS gdatatable (
                            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
                            db_name TEXT,
                            db_disease TEXT,
                            db_mobile TEXT,
                            db_age TEXT,
                            db_gender TEXT,
                            db_case TEXT,
                            db_record TEXT,
                            db_folder TEXT,
                            db_media TEXT
                        );"""

        
        self.cur.execute(data_table)

        
    
        #making tabs
        tabControl = ttk.Notebook(master)
        self.tab1 = ttk.Frame(tabControl)
        tabControl.add(self.tab1, text = "Record Book")
    
        self.tab2 = ttk.Frame(tabControl)
        tabControl.add(self.tab2, text = "Patient Registration")

        self.tab3 = ttk.Frame(tabControl)
        tabControl.add(self.tab3, text = "Update Record")

        tabControl.pack(expan = 1, fill = "both")

        
        
    
        #frames
        #tab2 frames
        self.top2 = Frame(self.tab2, height=150, bg='#99ff33')
        self.top2.pack(fill=X)

        self.bottom2 = Frame(self.tab2, height=550, bg='#00ffff')
        self.bottom2.pack(fill=X)

        #tab1 frames
        self.top1 = Frame(self.tab1, height=150, bg='#99ff33')
        self.top1.pack(fill=X)

        self.bottom1 = Frame(self.tab1, height=450, bg='#00ffff')
        self.bottom1.pack(fill= X)
        self.bottom1a = Frame(self.tab1, height=200, bg='#00ffff')
        self.bottom1a.pack(fill= X)

        #tab3 frames
        self.top3 = Frame(self.tab3, height=150, bg='#99ff33')
        self.top3.pack(fill=X)

        self.bottom3 = Frame(self.tab3, height=450, bg='#00ffff')
        self.bottom3.pack(fill= X)
        self.bottom3a = Frame(self.tab3, height=200, bg='#00ffff')
        self.bottom3a.pack(fill= X)


        #tab2 top frame design
        self.top_image = PhotoImage(file='icons/team.png')
        self.top_image_label = Label(self.top2, image=self.top_image, bg='#99ff33')
        self.top_image_label.place(x=65, y=12)

        self.heading = Label(self.top2, text='Add Patient Record ', font='sans 18 bold', bg='#99ff33', fg='#000066')
        self.heading.place(x=250, y=65)

        self.date_label = Label(self.top2, text="Today's date: " + date, font='arial 12 italic', bg='#99ff33', fg='#000066')
        self.date_label.place(x=650, y=12)

        #tab1 top frame design
        self.top_image1 = PhotoImage(file='icons/clipboard.png')
        self.top_image1_label = Label(self.top1, image=self.top_image1, bg='#99ff33')
        self.top_image1_label.place(x=65, y=12)

        self.heading1 = Label(self.top1, text='Patient Management App', font='sans 18 bold', bg='#99ff33', fg='#000066', cursor="hand2")
        self.heading1.place(x=250, y=65)

        
        self.heading1.bind("<Button-1>", lambda event: webbrowser.open("https://pmajm2help.blogspot.com"))

        self.date_label = Label(self.top1, text="Today's date: " + date, font='arial 12 italic', bg='#99ff33', fg='#000066')
        self.date_label.place(x=650, y=12)

        self.book_image = PhotoImage(file='icons/book.png')
        self.book_button = Button(self.top1, image= self.book_image, bg='#99ff33', command = self.book)
        self.book_button.place(x=750, y=90)

        self.prescription_image = PhotoImage(file='icons/prescription.png')
        self.prescription_button = Button(self.top1, image= self.prescription_image, bg='#99ff33', command = self.prescription)
        self.prescription_button.place(x=650, y=90)


        #tab2 bottom frame design
        
        self.bottom_image = PhotoImage(file='icons/medical.png')
        self.bottom_image_label = Label(self.bottom2, image=self.bottom_image, bg='#00ffff')
        self.bottom_image_label.place(x=325, y=40)

        
        self.bottom_i = PhotoImage(file='icons/submit.png')

        self.label1 = Label(self.bottom2, text = "Patient Name:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label1.place(x=45, y=25)
        self.textEdit1 = Entry(self.bottom2, width = 35, textvariable = self.name)
        self.textEdit1.place(x=200, y=27)

        self.label2 = Label(self.bottom2, text = "Disease:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label2.place(x=45, y=60)
        self.textEdit2 = Entry(self.bottom2, width = 35, textvariable = self.disease)
        self.textEdit2.place(x=200, y=62)

        self.label3 = Label(self.bottom2, text = "Age:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label3.place(x=45, y=95)
        self.textEdit3 = Entry(self.bottom2, width = 35, textvariable = self.age)
        self.textEdit3.place(x=200, y=97)

        self.label4 = Label(self.bottom2, text = "Mobile:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label4.place(x=45, y=130)
        self.textEdit4 = Entry(self.bottom2, width = 35, textvariable = self.mobile)
        self.textEdit4.place(x=200, y=132)

        self.label5 = Label(self.bottom2, text = "Gender:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label5.place(x=45, y=165)
        self.textEdit5 = Entry(self.bottom2, width = 35, textvariable = self.gender)
        self.textEdit5.place(x=200, y=167)

        #submit button
        self.button = Button(self.bottom2, image= self.bottom_i, bg='#fcb52f', command = self.submit)
        self.button.bind('<Return>', self.submit_enter)
        self.button.place(x=250, y=250)

        #tab1 bottom1 frame design
         
            #button and labels

        self.patient_image = PhotoImage(file='icons/user.png')
        self.disease_image = PhotoImage(file='icons/virus.png')
        self.mobile_image = PhotoImage(file='icons/smartphone.png')
        self.doctor_image = PhotoImage(file='icons/doctor.png')
        
        self.tab2patient_label = Label(self.bottom1, image=self.patient_image, bg='#00ffff')
        self.tab2patient_label.place(x=375, y=100)

        self.tab2disease_label = Label(self.bottom1, image=self.disease_image, bg='#00ffff')
        self.tab2disease_label.place(x=375, y=170)

        self.tab2mobile_label = Label(self.bottom1, image=self.mobile_image, bg='#00ffff')
        self.tab2mobile_label.place(x=375, y=240)
        

        self.tab2doctor_label = Label(self.bottom1a, image=self.doctor_image, bg='#00ffff')
        self.tab2doctor_label.place(x=690, y=0)

        self.tab2info = Label(self.bottom1, text = "Kindly select patient record from the list", font='arial 12 bold', bg='#00ffff', fg='#000000')
        self.tab2info.grid(row=0, column=2, padx=10, pady=10, sticky=N)
        

        self.media = Button(self.bottom1, text='Open Media Folder', width=18, command = self.media_function)
        self.media.place(x=680, y=55)

        self.opencase = Button(self.bottom1, text='Open Case History', width=18, command = self.opencase_function)
        self.opencase.place(x=515, y=55)

        self.openrecord = Button(self.bottom1, text='Open Record File', width=18, command = self.openrecord_function)
        self.openrecord.place(x=350, y=55)



           #Listbox and scrollbar design
        self.scroll = Scrollbar(self.bottom1, orient=VERTICAL)
        self.scrollh = Scrollbar(self.bottom1, orient=HORIZONTAL)
        self.listbox = Listbox(self.bottom1, width=45, height=18)
        self.listbox.bind('<<ListboxSelect>>', self.onselect)
        self.listbox.grid(row=0, column=0, padx=(10,0))
        self.scroll.grid(row=0, column=1, sticky=N+S)
        self.scrollh.grid(row=1, column=0, sticky=E+W)

        self.listbox.config(yscrollcommand=self.scroll.set)
        self.listbox.config(xscrollcommand=self.scrollh.set)
        self.scrollh.config(command=self.listbox.xview)
        self.scroll.config(command=self.listbox.yview)

        cursor = self.con.execute("select * from gdatatable order by id asc")
        persons= cursor.fetchall()
        
        count = 0
        for person in persons:

            self.listbox.insert(count, str(person[0])+". " + person[1])
            
            count +=1

        

        
        #tab1 bottom 1a frame design
        self.options = ["Patient Id", "   Name   ", "  Mobile   ", " Disease  "]
        self.clicked =StringVar()
        self.clicked.set(self.options[0])
        drop = OptionMenu(self.bottom1a,self.clicked,*self.options)
        drop.place(x=80, y=30)
        self.search = StringVar()
        self.searchtext = Entry(self.bottom1a, width = 20, textvariable = self.search)
        self.searchtext.bind('<Return>', self.onenter)
        self.searchtext.place(x=200, y=35)
        self.searchbtn = Button(self.bottom1a, text='Search', width=12, command = self.searchbtn_function)
        self.searchbtn.place(x=400, y=30)
        self.refreshbtn = Button(self.bottom1a, text='Refresh', width=12, command = self.refreshbtn_function)
        self.refreshbtn.place(x=400, y=75)

        #tab3 top frame design

        self.date_label = Label(self.top3, text="Today's date: " + date, font='arial 12 italic', bg='#99ff33', fg='#000066')
        self.date_label.place(x=650, y=12)

        self.top3_image = PhotoImage(file='icons/refresh.png')
        self.top3_image_label = Label(self.top3, image=self.top3_image, bg='#99ff33')
        self.top3_image_label.place(x=65, y=12)

        self.heading = Label(self.top3, text='Update Patient Record ', font='sans 18 bold', bg='#99ff33', fg='#000066')
        self.heading.place(x=250, y=65)


        #tab3 bottom frame design

        self.label3a = Label(self.bottom3, text = "Patient Name:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label3a.place(x=400, y=40)
        self.textEdit3a = Entry(self.bottom3, width = 35, textvariable = self.name_update)
        self.textEdit3a.place(x=555, y=42)

        self.label3b = Label(self.bottom3, text = "Disease:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label3b.place(x=400, y=75)
        self.textEdit3b = Entry(self.bottom3, width = 35, textvariable = self.disease_update)
        self.textEdit3b.place(x=555, y=77)

        self.label3c = Label(self.bottom3, text = "Age:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label3c.place(x=400, y=110)
        self.textEdit3c = Entry(self.bottom3, width = 35, textvariable = self.age_update)
        self.textEdit3c.place(x=555, y=112)

        self.label3d = Label(self.bottom3, text = "Mobile:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label3d.place(x=400, y=145)
        self.textEdit3d = Entry(self.bottom3, width = 35, textvariable = self.mobile_update)
        self.textEdit3d.place(x=555, y=147)

        self.label3e = Label(self.bottom3, text = "Gender:", font='arial 16', bg='#00ffff', fg='#000000')
        self.label3e.place(x=400, y=180)
        self.textEdit3e = Entry(self.bottom3, width = 35, textvariable = self.gender_update)
        self.textEdit3e.place(x=555, y=182)

        

        self.tab3info = Label(self.bottom3, text = "Kindly select patient record from the list", font='arial 12 bold', bg='#00ffff', fg='#000000')
        self.tab3info.grid(row=0, column=2, padx=10, pady=10, sticky=N)

        self.update = Button(self.bottom3, text='Update', width=15, command = self.update_function)
        self.update.place(x=580, y=250)

        #self.delete = Button(self.bottom3, text='Delete Permanently', width=12, command = self.delete_function)
        #self.delete.place(x=500, y=225)



        #Listbox and scrollbar design
        self.scroll3 = Scrollbar(self.bottom3, orient=VERTICAL)
        self.scrollh3 = Scrollbar(self.bottom3, orient=HORIZONTAL)
        self.listbox3 = Listbox(self.bottom3, width=45, height=18)
        self.listbox3.bind('<<ListboxSelect>>', self.onselect_update)
        self.listbox3.grid(row=0, column=0, padx=(10,0))
        self.scroll3.grid(row=0, column=1, sticky=N+S)
        self.scrollh3.grid(row=1, column=0, sticky=E+W)

        self.listbox3.config(yscrollcommand=self.scroll3.set)
        self.listbox3.config(xscrollcommand=self.scrollh3.set)
        self.scrollh3.config(command=self.listbox3.xview)
        self.scroll3.config(command=self.listbox3.yview)

        cursor3 = self.con.execute("select * from gdatatable order by id asc")
        persons3= cursor3.fetchall()
        
        count = 0
        for person3 in persons3:

            self.listbox3.insert(count, str(person3[0])+". " + person3[1])
            
            count +=1

        
        #tab3 bottom 3a frame design
        self.options3 = ["Patient Id", "   Name   ", "  Mobile   ", " Disease  "]
        self.clicked3 =StringVar()
        self.clicked3.set(self.options[0])
        drop3 = OptionMenu(self.bottom3a,self.clicked3,*self.options3)
        drop3.place(x=80, y=30)
        self.search3 = StringVar()
        self.searchtext3 = Entry(self.bottom3a, width = 20, textvariable = self.search3)
        self.searchtext3.bind('<Return>', self.onenter3)
        self.searchtext3.place(x=200, y=35)
        self.searchbtn3 = Button(self.bottom3a, text='Search', width=12, command = self.searchbtn3_function)
        self.searchbtn3.place(x=400, y=30)
        self.refreshbtn3 = Button(self.bottom3a, text='Refresh', width=12, command = self.refreshbtn3_function)
        self.refreshbtn3.place(x=400, y=75)

        self.bottom3a_image = PhotoImage(file='icons/steth.png')
        self.bottom3a_image_label = Label(self.bottom3a, image=self.bottom3a_image, bg='#00ffff')
        self.bottom3a_image_label.place(x=710, y=0)

        #declaring label for refreshing labels data on pressing different buttons
        self.tab2namedb = None

    def prescription(self):

        os.chdir(self.parent_dir)

        self.selected_item = self.listbox.curselection()

        if self.selected_item !=():

            
            self.record = self.listbox.get(self.selected_item)
            
            self.record_id = self.record.split(".")[0]
            query = "select * from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()
            

            self.qrecord = self.result[7]
            self.qid = self.result[0]
            self.qname = self.result[1]
            self.destination = ('prescription\prescription.xlsx')

            wb1 = openpyxl.load_workbook(self.qrecord)
            ws1 = wb1["Patient Record"]

            wb2 = openpyxl.load_workbook(self.destination) 
            ws2 = wb2["Prescription"]
            

            rowp = ws1.max_row
           #print(rowp)
            
            
            while rowp > 0:
                cells = ws1.cell(row = rowp, column = 4).value
                #print(cells)
                dummy=None
                if cells == dummy:
                    rowp -= 1
                else:
                    break
            #print(rowp)

            sourcemed = ws1.cell(row = rowp, column = 4)
            ws2.cell(20, 2).value = sourcemed.value
            ws2.cell(15, 3).value = self.qid
            ws2.cell(17, 3).value = self.qname
            ws2.cell(15, 9).value = date
            wb2.save(self.destination)
            os.startfile(self.destination)


        else:

            messagebox.showinfo("Warning", "Kindly select Patient from the list")



    def book(self):

        os.chdir(self.parent_dir)

    
        os.startfile(r'book\kentmm.pdf')


    def onenter(self, evt):

        if self.clicked.get() == "Patient Id":
            query = "select * from gdatatable where id = '{}'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked.get() == "   Name   ":
            query = "select * from gdatatable where db_name like '%{}%'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked.get() == "  Mobile   ":
            query = "select * from gdatatable where db_mobile like '%{}%'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked.get() == " Disease  ":
            query = "select * from gdatatable where db_disease like '%{}%'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

    def onenter3(self, evt):

        if self.clicked3.get() == "Patient Id":
            query = "select * from gdatatable where id = '{}'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked3.get() == "   Name   ":
            query = "select * from gdatatable where db_name like '%{}%'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked3.get() == "  Mobile   ":
            query = "select * from gdatatable where db_mobile like '%{}%'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked3.get() == " Disease  ":
            query = "select * from gdatatable where db_disease like '%{}%'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")


    def onselect_update(self, evt):

        self.selected_item = self.listbox3.curselection()

        if self.selected_item !=():



            
            self.record = self.listbox3.get(self.selected_item)
                    
            self.record_id = self.record.split(".")[0]
            query = "select * from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()
        
            self.name_update.set(self.result[1])
            self.disease_update.set(self.result[2])
            self.mobile_update.set(self.result[3])
            self.age_update.set(self.result[4])
            self.gender_update.set(self.result[5])

            
    def onselect(self, evt):

        self.selected_item = self.listbox.curselection()

        
        
        if self.tab2namedb:
            self.tab2namedb.destroy()
            self.tab2diseasedb.destroy()
            self.tab2mobiledb.destroy()
            
        

        if self.selected_item !=():

            self.qname = StringVar()
            self.qdisease = StringVar()
            self.qmobile = StringVar()


            
            self.record = self.listbox.get(self.selected_item)
                    
            self.record_id = self.record.split(".")[0]
            query = "select * from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()

                    
                    
            self.tab2namedb = Label(self.bottom1, textvariable =self.qname, font='arial 12 bold', fg='#000000', bg='#00ffff')
            self.tab2namedb.place(x=480, y=125)

            self.tab2diseasedb = Label(self.bottom1, textvariable =self.qdisease, font='arial 12 bold', fg='#000000', bg='#00ffff')
            self.tab2diseasedb.place(x=480, y=195)

            self.tab2mobiledb = Label(self.bottom1, textvariable =self.qmobile, font='arial 12 bold', fg='#000000', bg='#00ffff')
            self.tab2mobiledb.place(x=480, y=265)
                    
            self.qname.set(self.result[1])
            self.qdisease.set(self.result[2])
            self.qmobile.set(self.result[3])

    def update_function(self):

        self.selected_item = self.listbox3.curselection()



        if self.selected_item !=():

            
            self.record = self.listbox3.get(self.selected_item)
            
            self.record_id = self.record.split(".")[0]
            query = "select * from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()
            

            self.qrecord = self.result[7]
            self.qcase = self.result[6]
            self.qid = int(self.result[0])
            


            wb = openpyxl.load_workbook(self.qrecord)
            ws = wb["Patient Record"]
            ws['B1'] = self.name_update.get()
            ws['B2'] = self.disease_update.get()
            ws['B3'] = self.age_update.get()
            ws['D2'] = self.mobile_update.get()
            ws['D3'] = self.gender_update.get()
            
            wb.save(self.qrecord)

            #query_db_update = "UPDATE 'gdatatable` SET db_name = self.name_update.get(), 
            query_db_update = "UPDATE gdatatable SET db_name = ?, db_disease = ?, db_mobile = ?, db_age = ?, db_gender = ? WHERE id = ?"
            self.con.execute(query_db_update, (self.name_update.get(), self.disease_update.get(), self.mobile_update.get(), self.age_update.get(), self.gender_update.get(), self.qid))
            self.con.commit()


            cursor = self.con.execute("select * from gdatatable order by id asc")
            persons= cursor.fetchall()
            self.listbox.delete(0,END)
            self.listbox3.delete(0,END)
                
            count = 0
            for person in persons:

                self.listbox.insert(count, str(person[0])+". " + person[1])
                self.listbox3.insert(count, str(person[0])+". " + person[1])
            
                count +=1

            messagebox.showinfo("Success", "Record successfully updated")

        else:

            messagebox.showinfo("Warning", "Kindly select Patient from the list")


    def update_submit_function(self):

        return

    def searchbtn3_function(self):
            

        if self.clicked3.get() == "Patient Id":
            query = "select * from gdatatable where id = '{}'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked3.get() == "   Name   ":
            query = "select * from gdatatable where db_name like '%{}%'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked3.get() == "  Mobile   ":
            query = "select * from gdatatable where db_mobile like '%{}%'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked3.get() == " Disease  ":
            query = "select * from gdatatable where db_disease like '%{}%'".format(self.search3.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox3.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox3.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")



    def refreshbtn3_function(self):

        self.searchtext3.delete(0, END)
        self.searchtext3.update()
        self.clicked3.set(self.options3[0])

        cursor = self.con.execute("select * from gdatatable order by id asc")
        persons= cursor.fetchall()
        self.listbox3.delete(0,END)
                        
        count = 0
        for person in persons:

            self.listbox3.insert(count, str(person[0])+". " + person[1])
                    
            count +=1



 
   #refreshbtn defination
    def refreshbtn_function(self):

        self.searchtext.delete(0, END)
        self.searchtext.update()
        self.clicked.set(self.options[0])

        cursor = self.con.execute("select * from gdatatable order by id asc")
        persons= cursor.fetchall()
        self.listbox.delete(0,END)
                        
        count = 0
        for person in persons:

            self.listbox.insert(count, str(person[0])+". " + person[1])
                    
            count +=1

    #searchbtn defination
    def searchbtn_function(self):

        
        if self.clicked.get() == "Patient Id":
            query = "select * from gdatatable where id = '{}'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked.get() == "   Name   ":
            query = "select * from gdatatable where db_name like '%{}%'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked.get() == "  Mobile   ":
            query = "select * from gdatatable where db_mobile like '%{}%'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

        if self.clicked.get() == " Disease  ":
            query = "select * from gdatatable where db_disease like '%{}%'".format(self.search.get())
            persons = self.con.execute(query).fetchall()
            
            if persons !=[]:

                self.listbox.delete(0,END)
                            
                count = 0
                for person in persons:

                    self.listbox.insert(0, str(person[0])+". " + person[1] + " | "+ person[2] + " | "+ person[3])
                        
                    count +=1
            else:

                messagebox.showinfo("Warning", "No records found")

     
        


    #open media folder button defination
    def media_function(self):
        os.chdir(self.parent_dir)

        self.selected_item = self.listbox.curselection()

            
        

        if self.selected_item !=():



            
            self.record = self.listbox.get(self.selected_item)
                    
            self.record_id = self.record.split(".")[0]
            query = "select db_media from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()
            os.startfile(self.result[0])

        else:

            messagebox.showinfo("Warning", "Kindly select Patient from the list")

                 
    #delete button defiation
    def delete_function(self):
        os.chdir(self.parent_dir)

        self.selected_item = self.listbox3.curselection()

        if self.selected_item !=():


            asnwer = messagebox.askquestion("Delete", "Do you really want to Permanently delete the record")
            if asnwer == "yes":

                self.selected_item = self.listbox3.curselection()
                self.record = self.listbox3.get(self.selected_item)
                
                self.record_id = self.record.split(".")[0]
                query = "select db_folder from gdatatable where id = '{}'".format(self.record_id)
                self.result = self.con.execute(query).fetchone()
                self.qfile = self.result[0]

                
                query_delete = "DELETE from gdatatable where id = '{}'".format(self.record_id)
                self.cur.execute(query_delete)
                self.con.commit()

                shutil.rmtree(self.qfile)

                cursor = self.con.execute("select * from gdatatable order by id asc")
                persons= cursor.fetchall()
                self.listbox.delete(0,END)
                self.listbox3.delete(0,END)
                        
                count = 0
                for person in persons:

                    self.listbox.insert(count, str(person[0])+". " + person[1])
                    self.listbox3.insert(count, str(person[0])+". " + person[1])
                    
                    count +=1

                messagebox.showinfo("Success","Record deleted successfully")

        else:

            messagebox.showinfo("Warning", "Kindly select Patient from the list")
            
    #open record button function defination
    def openrecord_function(self):
        os.chdir(self.parent_dir)

        self.selected_item = self.listbox.curselection()



        if self.selected_item !=():

            
            self.record = self.listbox.get(self.selected_item)
            
            self.record_id = self.record.split(".")[0]
            query = "select db_record from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()
            

            self.qrecord = self.result[0]

            wb = openpyxl.load_workbook(self.qrecord)
            ws = wb["Patient Record"]
            
            row = ws.max_row
            if ws.cell(row, 1).value == date:

                os.startfile(self.qrecord)

               
            else:

                while row > 0:
                    cells = ws.cell(row = row, column = 1).value
                    print(cells)
                    dummy=None
                    if cells == dummy:
                        row -= 1
                    else:
                        break
                
                if ws.cell(row, 1).value == date:

                    os.startfile(self.qrecord)
                else:
                    last_row = row +1
                    ws.cell(last_row, 1).value = date
                    wb.save(self.qrecord)
                    
                    os.startfile(self.qrecord)

            
                print(row)


                
                

        else:

            messagebox.showinfo("Warning", "Kindly select Patient from the list")
            
    # Opencase button defination
    def opencase_function(self):

        os.chdir(self.parent_dir)
        
        self.selected_item = self.listbox.curselection()


        if self.selected_item !=():


            self.selected_item = self.listbox.curselection()
            self.record = self.listbox.get(self.selected_item)
            
            self.record_id = self.record.split(".")[0]
            query = "select db_case from gdatatable where id = '{}'".format(self.record_id)
            self.result = self.con.execute(query).fetchone()
            
            self.qcase = self.result[0]

            os.startfile(self.qcase)
                   
            

        else:

            messagebox.showinfo("Warning", "Kindly select Patient from the list")
    

    #submit button defination
    def submit(self):
        
        #creating folder

        #self.parent_dir = os.getcwd()
        
        
        
        if self.name.get() and self.disease.get() and self.mobile.get() !="":
            
            try:
                
                self.patientdata = os.path.join(self.parent_dir, 'Patient_Data')
                os.chdir(self.patientdata)

                query = "insert into 'gdatatable' (db_name, db_disease, db_mobile, db_age, db_gender) values(?,?,?,?,?)"
                self.cur.execute(query, (self.name.get(), self.disease.get(), self.mobile.get(), self.age.get(), self.gender.get()))
                self.con.commit()
                query_id_start = "UPDATE `gdatatable` SET `id`=1121 WHERE _rowid_='1'"
                self.con.execute(query_id_start)
                self.con.commit()

                #displaying patient id in messagebox
                idquery = "select id from gdatatable where db_name =? AND db_disease =? AND db_mobile =?"
                self.result = self.con.execute(idquery, (self.name.get(), self.disease.get(), self.mobile.get())).fetchone()
                self.printid = str(self.result[0])

                #assigning path variables

                self.folder = os.path.join('Patient_Data', self.printid)
                self.path = os.path.join(self.patientdata, self.printid)
                self.recordfilepath = os.path.join(self.folder, self.printid + '_record.xlsx')
                self.casefilepath = os.path.join(self.folder, self.printid + '_casetaking.docx')
                self.media = os.path.join(self.folder,self.printid + '_Media')

                query_db_update = "UPDATE gdatatable SET db_case = ?, db_record = ?, db_folder = ?, db_media = ? WHERE id = ?"
                self.con.execute(query_db_update, (self.casefilepath, self.recordfilepath, self.folder, self.media, self.printid))
                self.con.commit()

                
                
            
                os.mkdir(self.path)
            
                os.chdir(self.path)
                os.mkdir(self.printid + '_Media')
                
                #creating record book
                workbook = xlsxwriter.Workbook((self.printid) +'_record.xlsx')
                worksheet = workbook.add_worksheet('Patient Record')
                worksheet.freeze_panes(4, 0)
                wrap = workbook.add_format({'text_wrap':'true'})
                bold = workbook.add_format({'bold': 'true'})
                green = workbook.add_format({'bold': True, 'bg_color': '#ADFF2F', 'text_wrap':'true'})
                blue = workbook.add_format({'bold': True, 'bg_color': '#87CEEB', 'text_wrap':'true'})
                border = workbook.add_format({'border': 1, 'border_color':'#000000'})

                worksheet.write('A4', "Date")
                worksheet.write('B4', "Symptoms")
                worksheet.write('C4', "Rubric")
                worksheet.write('D4', "Medicine \n & Frequency")

                worksheet.write('E4', "Comments")
                worksheet.write('F4', "Audio\ \nVideo File")
                worksheet.write('G4', "Picture File")
                worksheet.write('H4', "Attachments 1")
                worksheet.write('I4', "Attachments 2")
                worksheet.write('A1', "Patient Name")
                worksheet.write('A2', "Diesease")
                worksheet.write('A3', "Age")
                worksheet.write('C1', "Patient Id")
                worksheet.write('C2', "Mobile")
                worksheet.write('C3', "Gender")
            
                #data from form
                worksheet.write('B1', self.name.get())
                worksheet.write('B2', self.disease.get())
                worksheet.write('B3', self.age.get())
                worksheet.write('D1', self.printid)
                worksheet.write('D2', self.mobile.get())
                worksheet.write('D3', self.gender.get())

                worksheet.set_column(0, 0, cell_format=wrap, width=13)
                worksheet.set_column(1, 2, cell_format=wrap, width=30)
                worksheet.set_column(3, 3, cell_format=wrap, width=20)
                worksheet.set_column(4, 4, cell_format=wrap, width=25)
                worksheet.set_column(5, 6, cell_format=wrap, width=12)
                worksheet.set_column(7, 8, cell_format=wrap, width=12)
                worksheet.set_row(0, None, cell_format=green)
                worksheet.set_row(1, None, cell_format=green)
                worksheet.set_row(2, None, cell_format=green)
                worksheet.set_row(3, None, cell_format=blue)


                workbook.close()


                

                #creating casehistory file

                document = Document()

                document.add_heading('Case History File', 0)

                run = document.add_paragraph().add_run('Patient Id = ' + self.printid)
        
                run.bold = True

                #document.add_paragraph('Patient Id = ' + self.printid)

                table = document.add_table(rows=3, cols=4)

                table.style = 'MediumGrid1-Accent5'

                

                #adding data in each cells
                wname = table.cell(0, 0)
                wname.text = 'Patient Name'

                wdisease = table.cell(1, 0)
                wdisease.text = 'Disease'

                wage = table.cell(2, 0)
                wage.text = 'Age'

                wmobile = table.cell(0, 2)
                wmobile.text = 'Mobile'

                wgender = table.cell(1, 2)
                wgender.text = 'Gender'
                

                wdate = table.cell(2, 2)
                wdate.text = 'Taken on Date'

                ename = table.cell(0, 1)
                ename.text = self.name.get()

                edisease = table.cell(1, 1)
                edisease.text = self.disease.get()

                eage = table.cell(2, 1)
                eage.text = self.age.get()

                emobile = table.cell(0, 3)
                emobile.text = self.mobile.get()

                egender = table.cell(1, 3)
                egender.text = self.gender.get()

                edate = table.cell(2, 3)
                edate.text = date

                document.add_paragraph('')
 

                document.add_paragraph(' ')


                document.save((self.printid) +'_casetaking.docx')

                cursor = self.con.execute("select * from gdatatable order by id asc")
                persons= cursor.fetchall()
                self.listbox.delete(0,END)
                self.listbox3.delete(0,END)
                
                count = 0
                for person in persons:

                    self.listbox.insert(count, str(person[0])+". " + person[1])
                    self.listbox3.insert(count, str(person[0])+". " + person[1])
            
                    count +=1
                

                              

                messagebox.showinfo("Sucess", "Record Added with Patient Id " + self.printid)

                self.textEdit1.delete(0, END)
                self.textEdit1.update()

                self.textEdit2.delete(0, END)
                self.textEdit2.update()  

                self.textEdit3.delete(0, END)
                self.textEdit3.update()

                self.textEdit4.delete(0, END)
                self.textEdit4.update()

                self.textEdit5.delete(0, END)
                self.textEdit5.update()
                

            except EXCEPTION as e:
                messagebox.showerror("Error", str(e))
        else:
            messagebox.showerror("Error", "Kindly fill all the fields", icon='warning')

        

    def submit_enter(self, evt):
            
            #creating folder

            #self.parent_dir = os.getcwd()
            
            
            
            if self.name.get() and self.disease.get() and self.mobile.get() !="":
                
                try:
                    
                    self.patientdata = os.path.join(self.parent_dir, 'Patient_Data')
                    os.chdir(self.patientdata)

                    query = "insert into 'gdatatable' (db_name, db_disease, db_mobile, db_age, db_gender) values(?,?,?,?,?)"
                    self.cur.execute(query, (self.name.get(), self.disease.get(), self.mobile.get(), self.age.get(), self.gender.get()))
                    self.con.commit()
                    query_id_start = "UPDATE `gdatatable` SET `id`=1121 WHERE _rowid_='1'"
                    self.con.execute(query_id_start)
                    self.con.commit()

                    #displaying patient id in messagebox
                    idquery = "select id from gdatatable where db_name =? AND db_disease =? AND db_mobile =?"
                    self.result = self.con.execute(idquery, (self.name.get(), self.disease.get(), self.mobile.get())).fetchone()
                    self.printid = str(self.result[0])

                    #assigning path variables

                    self.folder = os.path.join('Patient_Data', self.printid)
                    self.path = os.path.join(self.patientdata, self.printid)
                    self.recordfilepath = os.path.join(self.folder, self.printid + '_record.xlsx')
                    self.casefilepath = os.path.join(self.folder, self.printid + '_casetaking.docx')
                    self.media = os.path.join(self.folder,self.printid + '_Media')

                    query_db_update = "UPDATE gdatatable SET db_case = ?, db_record = ?, db_folder = ?, db_media = ? WHERE id = ?"
                    self.con.execute(query_db_update, (self.casefilepath, self.recordfilepath, self.folder, self.media, self.printid))
                    self.con.commit()

                    
                    
                
                    os.mkdir(self.path)
                
                    os.chdir(self.path)
                    os.mkdir(self.printid + '_Media')
                    
                    #creating record book
                    workbook = xlsxwriter.Workbook((self.printid) +'_record.xlsx')
                    worksheet = workbook.add_worksheet('Patient Record')
                    worksheet.freeze_panes(4, 0)
                    wrap = workbook.add_format({'text_wrap':'true'})
                    bold = workbook.add_format({'bold': 'true'})
                    green = workbook.add_format({'bold': True, 'bg_color': '#ADFF2F', 'text_wrap':'true'})
                    blue = workbook.add_format({'bold': True, 'bg_color': '#87CEEB', 'text_wrap':'true'})
                    border = workbook.add_format({'border': 1, 'border_color':'#000000'})

                    worksheet.write('A4', "Date")
                    worksheet.write('B4', "Symptoms")
                    worksheet.write('C4', "Rubric")
                    worksheet.write('D4', "Medicine \n & Frequency")

                    worksheet.write('E4', "Comments")
                    worksheet.write('F4', "Audio\ \nVideo File")
                    worksheet.write('G4', "Picture File")
                    worksheet.write('H4', "Attachments 1")
                    worksheet.write('I4', "Attachments 2")
                    worksheet.write('A1', "Patient Name")
                    worksheet.write('A2', "Diesease")
                    worksheet.write('A3', "Age")
                    worksheet.write('C1', "Patient Id")
                    worksheet.write('C2', "Mobile")
                    worksheet.write('C3', "Gender")
                
                    #data from form
                    worksheet.write('B1', self.name.get())
                    worksheet.write('B2', self.disease.get())
                    worksheet.write('B3', self.age.get())
                    worksheet.write('D1', self.printid)
                    worksheet.write('D2', self.mobile.get())
                    worksheet.write('D3', self.gender.get())

                    worksheet.set_column(0, 0, cell_format=wrap, width=13)
                    worksheet.set_column(1, 2, cell_format=wrap, width=30)
                    worksheet.set_column(3, 3, cell_format=wrap, width=20)
                    worksheet.set_column(4, 4, cell_format=wrap, width=25)
                    worksheet.set_column(5, 6, cell_format=wrap, width=12)
                    worksheet.set_column(7, 8, cell_format=wrap, width=12)
                    worksheet.set_row(0, None, cell_format=green)
                    worksheet.set_row(1, None, cell_format=green)
                    worksheet.set_row(2, None, cell_format=green)
                    worksheet.set_row(3, None, cell_format=blue)


                    workbook.close()


                    

                    #creating casehistory file

                    document = Document()

                    document.add_heading('Case History File', 0)

                    run = document.add_paragraph().add_run('Patient Id = ' + self.printid)
            
                    run.bold = True

                    #document.add_paragraph('Patient Id = ' + self.printid)

                    table = document.add_table(rows=3, cols=4)

                    table.style = 'MediumGrid1-Accent5'

                    

                    #adding data in each cells
                    wname = table.cell(0, 0)
                    wname.text = 'Patient Name'

                    wdisease = table.cell(1, 0)
                    wdisease.text = 'Disease'

                    wage = table.cell(2, 0)
                    wage.text = 'Age'

                    wmobile = table.cell(0, 2)
                    wmobile.text = 'Mobile'

                    wgender = table.cell(1, 2)
                    wgender.text = 'Gender'
                    

                    wdate = table.cell(2, 2)
                    wdate.text = 'Taken on Date'

                    ename = table.cell(0, 1)
                    ename.text = self.name.get()

                    edisease = table.cell(1, 1)
                    edisease.text = self.disease.get()

                    eage = table.cell(2, 1)
                    eage.text = self.age.get()

                    emobile = table.cell(0, 3)
                    emobile.text = self.mobile.get()

                    egender = table.cell(1, 3)
                    egender.text = self.gender.get()

                    edate = table.cell(2, 3)
                    edate.text = date

                    document.add_paragraph('')
    

                    document.add_paragraph(' ')


                    document.save((self.printid) +'_casetaking.docx')

                    cursor = self.con.execute("select * from gdatatable order by id asc")
                    persons= cursor.fetchall()
                    self.listbox.delete(0,END)
                    self.listbox3.delete(0,END)
                    
                    count = 0
                    for person in persons:

                        self.listbox.insert(count, str(person[0])+". " + person[1])
                        self.listbox3.insert(count, str(person[0])+". " + person[1])
                
                        count +=1
                    

                                

                    messagebox.showinfo("Sucess", "Record Added with Patient Id " + self.printid)

                    self.textEdit1.delete(0, END)
                    self.textEdit1.update()

                    self.textEdit2.delete(0, END)
                    self.textEdit2.update()  

                    self.textEdit3.delete(0, END)
                    self.textEdit3.update()

                    self.textEdit4.delete(0, END)
                    self.textEdit4.update()

                    self.textEdit5.delete(0, END)
                    self.textEdit5.update()
                    

                except EXCEPTION as e:
                    messagebox.showerror("Error", str(e))
            else:
                messagebox.showerror("Error", "Kindly fill all the fields", icon='warning')

            

def main():
    root = Tk()
    app = Application(root)

    root.title("Patient Management App JM2 by gspc")
    root.geometry('850x600')
    #root.resizable(False,False)
    root.iconbitmap('icons\icon.ico')
    
    
    root.mainloop()


if __name__=='__main__':
    main()